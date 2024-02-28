from django.shortcuts import render
from django.http import HttpResponse
from PlagiarismAnalyzerApp.algorithm import main
from docx import *
from PlagiarismAnalyzerApp.algorithm import fileSimilarity
import PyPDF2 
from django.shortcuts import render
import nltk
from nltk.corpus import wordnet
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords 
import random

stop_words = stopwords.words("english") 

def plagiarism_remover(i):
    word = i
    synonyms = []
    if word in stop_words:
        return word
    if wordnet.synsets(word)==[]:
        return word
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            synonyms.append(lemma.name())
    pos_tag_word = nltk.pos_tag([word])
    pos = []
    for i in synonyms:
        pos.append(nltk.pos_tag([i]))
    final_synonyms = []
    for i in pos:
        if pos_tag_word[0][1] == i[0][1]:
            final_synonyms.append(i[0][0])
    final_synonyms = list(set(final_synonyms))
    if final_synonyms == []:
        return word
    if word.istitle():
        return random.choice(final_synonyms).title()
    else:
        return random.choice(final_synonyms)
    

def plagiarism_removal(para):
    para_split = word_tokenize(para)
    final_text = []
    for i in para_split:
        final_text.append(plagiarism_remover(i))
    final_text = " ".join(final_text)
    return final_text


def plrem(request):
    if request.method == 'POST':
        paragraph = request.POST['paragraph']
        result = plagiarism_removal(paragraph)
        return render(request, 'pc/plrem.html', {'result': result})
    else:
        return render(request, 'pc/plrem.html')

    
#about

def about(request):
    return render(request, 'pc/about.html')

# Create your views here.
#home
def home(request):
    return render(request, 'pc/index.html') 

#web search(Text)
def test(request):
    print("request is welcome test")
    #print(request.POST['q'])  
    
    if request.POST['q']: 
        percent,link = main.findSimilarity(request.POST['q'])
        percent = round(percent,2)
    print("Output.....................!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': int(percent)})

#web search file(.txt, .docx)
def filetest(request):
    value = ''    
    print(request.FILES['docfile'])
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read())

    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text

    elif str(request.FILES['docfile']).endswith(".pdf"):
        # creating a pdf file object 
        pdfFileObj = open(request.FILES['docfile'], 'rb') 

        # creating a pdf reader object 
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 

        # printing number of pages in pdf file 
        print(pdfReader.numPages) 

        # creating a page object 
        pageObj = pdfReader.getPage(0) 

        # extracting text from page 
        print(pageObj.extractText()) 

        # closing the pdf file object 
        pdfFileObj.close() 


    percent,link = main.findSimilarity(value)
    print("Output...................!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': int(percent)})

#text compare
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

#two text compare(Text)
def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
    result = round(result,2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': int(result)})
    

#two text compare(.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1,value2)
    
    print("Output..................!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': int(result)})



